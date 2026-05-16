"""Общие утилиты: нормализация, конфиг."""

import os
import re
import sys
import json


def app_dir():
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(sys.argv[0]))


def res_path(name):
    """Путь рядом с exe или скриптом."""
    return os.path.join(app_dir(), name)


def load_config():
    p = res_path("config.json")
    if os.path.exists(p):
        try:
            with open(p, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {}


def save_config(data):
    p = res_path("config.json")
    try:
        with open(p, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def load_settings():
    p = res_path("settings.json")
    if os.path.exists(p):
        try:
            with open(p, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {}


def save_settings(data):
    p = res_path("settings.json")
    try:
        with open(p, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


# ---- нормализация для поиска ----

def normalize_phone(s):
    digits = re.sub(r"\D", "", str(s or ""))
    if not digits:
        return ""
    if len(digits) == 11 and digits[0] in ("7", "8"):
        digits = digits[1:]
    if len(digits) >= 10:
        return digits[-10:]
    return digits


def normalize_fio(s):
    s = str(s or "").lower().replace("ё", "е").replace("-", " ")
    return re.sub(r"\s+", " ", s).strip()


def is_phone_query(s):
    return len(re.sub(r"\D", "", str(s or ""))) >= 10


def fmt_phone(d10):
    return "+7" + d10 if d10 and len(d10) == 10 else (d10 or "")


def fmt_date(s):
    s = str(s or "")
    m = re.match(r"(\d{4})-(\d{2})-(\d{2})", s)
    if m:
        y, mo, d = m.groups()
        return f"{d}.{mo}.{y}"
    return s


# ---- извлечение из произвольных файлов ----

def extract_text_from_docx(path):
    """Возвращает плоский текст из docx (включая таблицы)."""
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def find_phones_in_text(text):
    """Все номера в тексте, нормализованные до 10 цифр."""
    found = set()
    for m in re.finditer(r"(?:\+?7|8)?[\s\-\(\)]?\d{3}[\s\-\(\)]?\d{3}[\s\-\(\)]?\d{2}[\s\-\(\)]?\d{2}", text):
        d = normalize_phone(m.group(0))
        if len(d) == 10:
            found.add(d)
    return found


def find_fios_in_text(text):
    """Грубый поиск ФИО в тексте: три слова с заглавных букв подряд."""
    pattern = re.compile(r"\b([А-ЯЁ][а-яё]+)\s+([А-ЯЁ][а-яё]+)\s+([А-ЯЁ][а-яё]+)\b")
    return [(m.group(1), m.group(2), m.group(3)) for m in pattern.finditer(text)]
